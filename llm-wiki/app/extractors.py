from __future__ import annotations

import hashlib
import logging
from pathlib import Path
import re

log = logging.getLogger(__name__)

# Bildendungen -> MIME-Typ fuer die data-URL an das Sprachmodell.
IMAGE_EXTENSIONS = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".heic": "image/heic",
}

LEGACY_DOC_HINT = (
    "Das alte Word-Format (.doc) kann nicht gelesen werden. "
    "Bitte als DOCX speichern und erneut hochladen."
)
LEGACY_XLS_HINT = (
    "Das alte Excel-Format (.xls) kann nicht gelesen werden. "
    "Bitte als XLSX speichern und erneut hochladen."
)


def extract_pdf(file_path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages_text = []
    for idx, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages_text.append(f"<!-- Seite {idx} -->\n{text}")
    return "\n\n".join(pages_text)


def extract_docx(file_path: Path) -> str:
    import docx

    doc = docx.Document(str(file_path))
    content_blocks = []

    # 1. Absaetze & Ueberschriften
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style_name = p.style.name.lower() if p.style else ""
        if "heading 1" in style_name:
            content_blocks.append(f"# {text}")
        elif "heading 2" in style_name:
            content_blocks.append(f"## {text}")
        elif "heading 3" in style_name:
            content_blocks.append(f"### {text}")
        else:
            content_blocks.append(text)

    # 2. Tabellen in sauberes Markdown-Format umwandeln
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(row_cells):
                rows_data.append(row_cells)
        if rows_data:
            header = rows_data[0]
            col_count = max(len(header), max((len(r) for r in rows_data), default=1))
            padded_header = header + [""] * (col_count - len(header))
            md_table = [
                "| " + " | ".join(padded_header) + " |",
                "| " + " | ".join(["---"] * col_count) + " |",
            ]
            for row in rows_data[1:]:
                padded_row = row + [""] * (col_count - len(row))
                md_table.append("| " + " | ".join(padded_row[:col_count]) + " |")
            content_blocks.append("\n".join(md_table))

    return "\n\n".join(content_blocks)


def extract_xlsx(file_path: Path) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(str(file_path), data_only=True, read_only=True)
    sheets_output = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        # Leere Zeilen filtern
        non_empty_rows = [
            r for r in rows if any(cell is not None and str(cell).strip() != "" for cell in r)
        ]
        if not non_empty_rows:
            continue

        sheet_md = [f"## Tabellenblatt: {sheet_name}\n"]
        max_cols = max(len(r) for r in non_empty_rows)

        # Erste nicht-leere Zeile als Header
        header = [str(c) if c is not None else "" for c in non_empty_rows[0]]
        header += [""] * (max_cols - len(header))
        sheet_md.append("| " + " | ".join(header) + " |")
        sheet_md.append("| " + " | ".join(["---"] * max_cols) + " |")

        for r in non_empty_rows[1:]:
            row_vals = [str(c).replace("\n", " ").strip() if c is not None else "" for c in r]
            row_vals += [""] * (max_cols - len(row_vals))
            sheet_md.append("| " + " | ".join(row_vals[:max_cols]) + " |")

        sheets_output.append("\n".join(sheet_md))

    wb.close()
    return "\n\n".join(sheets_output)


# Beschreibungen je Datei-Hash: das Drop-In (Prefill) und der Upload sehen
# dieselben Bytes - das Modell soll dafuer nur einmal laufen.
_IMAGE_CACHE: dict[str, str] = {}
_IMAGE_CACHE_MAX = 32


# Antworten, in denen das Modell sagt, dass kein Bild angekommen ist (z.B. wenn
# ein Proxy `image_url` verschluckt). Das ist keine Beschreibung, sondern ein
# Fehlschlag - dann lieber der ehrliche Fallback mit den Bildmassen.
_NO_IMAGE_RE = re.compile(
    r"(kann (das|dieses|ein) bild\b[^.]{0,40}?nicht (sehen|erkennen|anzeigen|öffnen|oeffnen)"
    r"|keine? bild(datei|daten|inhalt)"
    r"|(wurde|ist) (mir )?(nicht|kein(e)?)\s?(übermittelt|uebermittelt|angehängt|angehaengt|mitgeschickt)"
    r"|kein bild (angehängt|angehaengt|übermittelt|uebermittelt|vorhanden|erhalten)"
    r"|(can(no|')t|cannot|unable to) (see|view|access) (the|any|this) image"
    r"|no image (was |has been )?(provided|attached|received|uploaded))",
    re.IGNORECASE,
)


def looks_like_no_image(text: str) -> bool:
    """True, wenn die Modellantwort nur sagt, dass kein Bild uebermittelt wurde."""
    return bool(_NO_IMAGE_RE.search(text[:400]))


def image_size(file_path: Path) -> tuple[int, int] | None:
    """(Breite, Hoehe) per Pillow; None, wenn Pillow das Format nicht kennt (z.B. HEIC ohne Plugin)."""
    try:
        from PIL import Image

        with Image.open(str(file_path)) as im:
            return im.size
    except Exception:
        return None


def image_fallback_text(filename: str, size: tuple[int, int] | None) -> str:
    masse = f"{size[0]}\u00d7{size[1]} px" if size else "Masse unbekannt"
    return (
        f"Bilddatei {filename}, {masse}. "
        "Kein Text extrahiert (Sprachmodell nicht erreichbar)."
    )


def extract_image(file_path: Path, filename: str) -> str:
    """Bild -> Beschreibung durch das Sprachmodell. Ohne Key, bei Fehler oder
    Timeout kommt der Fallback-Text mit den Bildmassen - nie eine Exception,
    der Upload muss auch ohne Modell durchlaufen."""
    ext = Path(filename).suffix.lower()
    mime = IMAGE_EXTENSIONS.get(ext, "application/octet-stream")
    data = file_path.read_bytes()
    key = hashlib.sha256(data).hexdigest()
    text = _IMAGE_CACHE.get(key, "")
    if not text:
        try:
            from . import llm

            text = (llm.describe_image(data, mime, filename) or "").strip()
            if text and looks_like_no_image(text):
                log.warning("Bildbeschreibung fuer %s: Modell hat kein Bild erhalten (%.120r)",
                            filename, text)
                text = ""
        except Exception as exc:
            log.warning("Bildbeschreibung fuer %s fehlgeschlagen: %s", filename, exc)
            text = ""
    if text:
        if len(_IMAGE_CACHE) >= _IMAGE_CACHE_MAX:
            _IMAGE_CACHE.pop(next(iter(_IMAGE_CACHE)))
        _IMAGE_CACHE[key] = text
        return text
    return image_fallback_text(filename, image_size(file_path))


def extract_text_file(file_path: Path) -> str:
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="latin-1", errors="replace")


# Bilder, die per Upload als Original angehaengt werden (HEIC bewusst nicht:
# Browser und Pillow koennen es ohne Plugin nicht anzeigen).
UPLOAD_IMAGE_SUFFIXES = (".jpg", ".jpeg", ".png", ".gif", ".webp")


def is_image(filename: str) -> bool:
    return Path(filename).suffix.lower() in UPLOAD_IMAGE_SUFFIXES


def extract_text_from_file(file_path: Path, filename: str) -> str:
    """Extrahiert Volltext aus einer hochgeladenen Datei je nach Dateiendung.

    Bilder werfen nie: mit Modell liefert extract_image eine Beschreibung,
    ohne Modell (oder bei Fehler/Timeout) den Fallback-Text mit den Bildmassen.
    """
    ext = Path(filename).suffix.lower()

    if ext in IMAGE_EXTENSIONS:
        # Nie eine Exception: ohne Modell (oder bei Fehler/Timeout) kommt der
        # Fallback-Text mit den Bildmassen, der Upload laeuft trotzdem durch.
        return extract_image(file_path, filename)
    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext == ".docx":
        return extract_docx(file_path)
    elif ext == ".xlsx":
        return extract_xlsx(file_path)
    elif ext == ".doc":
        # python-docx liest nur OOXML; ein echtes .doc endet sonst im Traceback.
        try:
            return extract_docx(file_path)
        except Exception as exc:
            raise ValueError(LEGACY_DOC_HINT) from exc
    elif ext == ".xls":
        try:
            return extract_xlsx(file_path)
        except Exception as exc:
            raise ValueError(LEGACY_XLS_HINT) from exc
    elif ext in (".md", ".markdown", ".txt"):
        return extract_text_file(file_path)
    else:
        # Fallback: Versuche als Textdatei einzulesen
        return extract_text_file(file_path)
